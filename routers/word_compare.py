import io
import logging
from fastapi import APIRouter, UploadFile, File, HTTPException
from docx import Document
from difflib import SequenceMatcher

logger = logging.getLogger(__name__)

router = APIRouter()

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        file_like = io.BytesIO(file_content)
        doc = Document(file_like)
        
        # Extract all paragraphs (keep empty ones for structure)
        text_lines = []
        for paragraph in doc.paragraphs:
            text_lines.append(paragraph.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                if any(t.strip() for t in row_text):
                    text_lines.append(" | ".join(row_text))
        
        return "\n".join(text_lines)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Error reading DOCX file: {str(e)}")

def get_inline_diff(old_line: str, new_line: str) -> tuple:
    """
    Character-level diff between two lines.
    Returns (old_segments, new_segments) where each segment is {'text': str, 'type': 'equal'|'removed'|'added'}
    Spaces are preserved and highlighted when they differ.
    """
    char_matcher = SequenceMatcher(None, old_line, new_line)
    old_segments = []
    new_segments = []
    
    for tag, i1, i2, j1, j2 in char_matcher.get_opcodes():
        if tag == 'equal':
            old_segments.append({'text': old_line[i1:i2], 'type': 'equal'})
            new_segments.append({'text': new_line[j1:j2], 'type': 'equal'})
        elif tag == 'delete':
            old_segments.append({'text': old_line[i1:i2], 'type': 'removed'})
        elif tag == 'insert':
            new_segments.append({'text': new_line[j1:j2], 'type': 'added'})
        elif tag == 'replace':
            old_segments.append({'text': old_line[i1:i2], 'type': 'removed'})
            new_segments.append({'text': new_line[j1:j2], 'type': 'added'})
    
    return old_segments, new_segments


def compare_texts(text1: str, text2: str) -> dict:
    """Compare two texts and return differences"""
    lines1 = text1.split('\n')
    lines2 = text2.split('\n')
    
    # Use SequenceMatcher to get matching blocks
    matcher = SequenceMatcher(None, lines1, lines2)
    
    file1_result = []
    file2_result = []
    unified_result = []
    
    # Process all match operations
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            # Same content in both files
            for line in lines1[i1:i2]:
                file1_result.append({
                    'type': 'unchanged',
                    'content': line
                })
            for line in lines2[j1:j2]:
                file2_result.append({
                    'type': 'unchanged',
                    'content': line
                })
                unified_result.append({
                    'type': 'unchanged',
                    'content': line
                })
                
        elif tag == 'delete':
            # Only in file1 (removed)
            for line in lines1[i1:i2]:
                file1_result.append({
                    'type': 'removed',
                    'content': line
                })
                unified_result.append({
                    'type': 'removed',
                    'content': line
                })
                
        elif tag == 'insert':
            # Only in file2 (added)
            for line in lines2[j1:j2]:
                file2_result.append({
                    'type': 'added',
                    'content': line
                })
                unified_result.append({
                    'type': 'added',
                    'content': line
                })
                
        elif tag == 'replace':
            # Character-level inline diff for paired lines
            removed_lines = lines1[i1:i2]
            added_lines = lines2[j1:j2]
            
            # Pair lines up for inline diff
            paired = min(len(removed_lines), len(added_lines))
            
            for k in range(paired):
                old_segments, new_segments = get_inline_diff(removed_lines[k], added_lines[k])
                file1_result.append({
                    'type': 'removed',
                    'content': removed_lines[k],
                    'segments': old_segments
                })
                file2_result.append({
                    'type': 'added',
                    'content': added_lines[k],
                    'segments': new_segments
                })
                unified_result.append({
                    'type': 'removed',
                    'content': removed_lines[k],
                    'segments': old_segments
                })
                unified_result.append({
                    'type': 'added',
                    'content': added_lines[k],
                    'segments': new_segments
                })
            
            # Remaining unpaired lines (no inline diff)
            for line in removed_lines[paired:]:
                file1_result.append({
                    'type': 'removed',
                    'content': line
                })
                unified_result.append({
                    'type': 'removed',
                    'content': line
                })
            for line in added_lines[paired:]:
                file2_result.append({
                    'type': 'added',
                    'content': line
                })
                unified_result.append({
                    'type': 'added',
                    'content': line
                })
    
    return {
        'file1': file1_result,
        'file2': file2_result,
        'unified': unified_result
    }

@router.post('/word-compare')
async def word_compare(file1: UploadFile = File(...), file2: UploadFile = File(...)):
    """Compare two Word documents"""
    try:
        # Validate file types
        if not file1.filename.endswith('.docx') or not file2.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail='Only .docx files are supported')
        
        # Read files
        file1_content = await file1.read()
        file2_content = await file2.read()
        
        # Extract text
        logger.info(f"Extracting text from {file1.filename}")
        text1 = extract_text_from_docx(file1_content)
        
        logger.info(f"Extracting text from {file2.filename}")
        text2 = extract_text_from_docx(file2_content)
        
        # Compare
        logger.info("Comparing documents")
        result = compare_texts(text1, text2)
        
        logger.info("Document comparison completed successfully")
        return result
        
    except ValueError as e:
        logger.error(f"Validation error: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error comparing documents: {e}")
        raise HTTPException(status_code=500, detail='Error processing documents')
